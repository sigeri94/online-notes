from lxml import etree
import argparse
import pandas as pd
import socket
import struct
from datetime import datetime
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

def parse_nessus_file_for_vulns(nessus_file, vuln_name_filters):
    vulnerabilities = []

    # Parse the .nessus file
    tree = etree.parse(nessus_file)

    # List of keywords to mark as skipped
    skipped_keywords = ["ms09-050", "redis"]

    # Extract hosts and vulnerabilities
    for report_host in tree.xpath("//ReportHost"):
        ip_address = report_host.attrib.get("name")

        for report_item in report_host.xpath(".//ReportItem"):
            vuln_name = report_item.attrib.get("pluginName", "Unknown Vulnerability")
            severity_code = report_item.attrib.get("severity", "0")
            solution = report_item.findtext("solution", "No solution provided")
            plugin_output = report_item.findtext("plugin_output", "No plugin output provided")

            # Convert severity code to readable format
            severity_mapping = {
                "0": "Informational",
                "1": "Low",
                "2": "Medium",
                "3": "High",
                "4": "Critical"
            }
            severity = severity_mapping.get(severity_code, "Unknown")

            # Check if the plugin output should be marked as skipped
            if any(keyword in vuln_name.lower() for keyword in skipped_keywords):
                plugin_output = "Skipped"

            # Check if the vulnerability name matches any of the filters
            if any(filter_name.lower() in vuln_name.lower() for filter_name in vuln_name_filters):
                vulnerabilities.append((ip_address, vuln_name, severity, plugin_output, solution))

    return vulnerabilities

def add_cover_page(doc, report_date):
    """Add a cover page to the DOCX document."""
    cover_title = "Vulnerability Report"
    description = "This document provides a detailed analysis of vulnerabilities identified during a penetration test."

    # Add title
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title_paragraph.add_run(cover_title)
    run.bold = True
    run.font.size = Pt(24)

    # Add empty lines
    doc.add_paragraph()
    doc.add_paragraph()

    # Add date
    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_paragraph.add_run(f"Date: {report_date}").font.size = Pt(14)

    # Add empty lines
    doc.add_paragraph()
    doc.add_paragraph()

    # Add description
    description_paragraph = doc.add_paragraph()
    description_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    description_paragraph.add_run(description).font.size = Pt(12)

    # Add page break
    doc.add_page_break()

from collections import defaultdict

def save_to_docx(vulnerabilities, output_filename):
    # Create a new DOCX document
    doc = Document()

    # Generate report date
    report_date = datetime.now().strftime("%B %d, %Y")

    # Add cover page
    add_cover_page(doc, report_date)

    # Add title
    title = doc.add_heading("Vulnerability Report", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add Table of Contents
    doc.add_heading("Table of Contents", level=2)
    doc.add_paragraph("1. Scope", style="List Number")
    doc.add_paragraph("2. Summary of Findings", style="List Number")
    doc.add_paragraph("3. Details of Findings", style="List Number")

    # Add Scope
    doc.add_heading("1. Scope", level=2)
    doc.add_paragraph("This report summarizes the vulnerabilities identified during a penetration test of the target systems.")

    # Add Summary of Findings
    doc.add_heading("2. Summary of Findings", level=2)

    # Create summary
    df = pd.DataFrame(vulnerabilities, columns=["IP Address", "Vulnerability Name", "Severity", "Plugin Output", "Solution"])
    severity_counts = df["Severity"].value_counts().to_dict()

    # Add severity table
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Severity"
    hdr_cells[1].text = "Count"

    for severity, count in severity_counts.items():
        row_cells = table.add_row().cells
        row_cells[0].text = severity
        row_cells[1].text = str(count)

    # Add Details of Findings
    doc.add_heading("3. Details of Findings", level=2)

    # Group vulnerabilities by Vulnerability Name
    findings_by_vuln_name = defaultdict(list)
    for ip_address, vuln_name, severity, plugin_output, solution in vulnerabilities:
        findings_by_vuln_name[vuln_name].append((ip_address, severity, plugin_output, solution))

    # Add findings grouped by Vulnerability Name
    for vuln_name, findings in sorted(findings_by_vuln_name.items()):
        # Add Vulnerability Name Header
        doc.add_heading(f"Vulnerability Name: {vuln_name}", level=3)

        # Get unique severity and solution (assuming they're consistent across all IPs for a given vuln)
        severity = findings[0][1]
        solution = findings[0][3]

        # Add Severity
        doc.add_paragraph(f"Severity: {severity}", style="Normal")

        # List affected IP Addresses
        doc.add_paragraph("Affected IP Addresses:", style="Normal")
        for ip_address, _, _, _ in findings:
            doc.add_paragraph(f"- {ip_address}", style="List Bullet")

        # Add Solution
        doc.add_paragraph("Solution:", style="Normal")
        doc.add_paragraph(solution, style="Normal")

        # Add separator
        doc.add_paragraph("---", style="Normal")

    # Save the document
    doc.save(output_filename)
    print(f"Saved DOCX file to: {output_filename}")

def main():
    # Set up argument parser
    parser = argparse.ArgumentParser(description="Process a .nessus file and extract specific vulnerabilities.")
    parser.add_argument("nessus_file", help="Path to the .nessus file to be processed")
    parser.add_argument("vuln_name_filters", nargs="+", help="Names of the vulnerabilities to filter (multiple allowed)")
    args = parser.parse_args()

    # Parse the .nessus file for the specified vulnerability names
    vulnerabilities = parse_nessus_file_for_vulns(args.nessus_file, args.vuln_name_filters)

    if not vulnerabilities:
        print("No matching vulnerabilities found.")
        return

    # Generate timestamp for filenames
    timestamp = datetime.now().strftime("%Y%m%d%H%M")

    # Define filename
    docx_filename = f"vulnerability_report_{timestamp}.docx"

    # Save results to DOCX
    save_to_docx(vulnerabilities, docx_filename)

if __name__ == '__main__':
    main()
