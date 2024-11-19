from lxml import etree
import argparse
import pandas as pd
from collections import defaultdict
from datetime import datetime
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement


def create_shaded_table(doc, content, shading_color="D9D9D9"):
    """Create a shaded table for plugin output."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    cell.text = content

    # Apply shading to the table cell using a different method
    cell_xml = cell._element
    shading = OxmlElement("w:shd")
    #shading.set("w:fill", shading_color)  # Set the shading color (e.g., D9D9D9 for light gray)
    cell_xml.get_or_add_tcPr().append(shading)

    # Style the text within the cell
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 0, 0)

    return table


def parse_vuln_file(vuln_file):
    """Parse the vuln.txt file to extract the list of vulnerability names."""
    try:
        with open(vuln_file, "r") as file:
            return [line.strip() for line in file if line.strip()]
    except FileNotFoundError:
        print(f"Error: File '{vuln_file}' not found.")
        return []


def parse_ip_file(ip_file):
    """Parse the IP file to extract hostname, internal IP, and public IP."""
    ip_data = {}
    try:
        with open(ip_file, "r") as file:
            for line in file:
                parts = line.strip().split()
                if len(parts) >= 3:
                    hostname, internal_ip, public_ip = parts[0], parts[1], parts[2]
                    ip_data[internal_ip] = {"hostname": hostname, "public_ip": public_ip}
        return ip_data
    except FileNotFoundError:
        print(f"Error: File '{ip_file}' not found.")
        return {}


def parse_nessus_file_for_vulns(nessus_file, vuln_filters, ip_filter, ip_metadata):
    vulnerabilities = []

    # Parse the .nessus file
    tree = etree.parse(nessus_file)

    # Extract hosts and vulnerabilities
    for report_host in tree.xpath("//ReportHost"):
        ip_address = report_host.attrib.get("name")

        # Skip IPs not in the filter
        if ip_filter and ip_address not in ip_filter:
            continue

        for report_item in report_host.xpath(".//ReportItem"):
            vuln_name = report_item.attrib.get("pluginName", "Unknown Vulnerability")
            severity_code = report_item.attrib.get("severity", "0")
            solution = report_item.findtext("solution", "No solution provided")
            plugin_output = report_item.findtext("plugin_output", "No plugin output provided")
            description = report_item.findtext("description", "No description available")

            # Convert severity code to readable format
            severity_mapping = {
                "0": "Informational",
                "1": "Low",
                "2": "Medium",
                "3": "High",
                "4": "Critical"
            }
            severity = severity_mapping.get(severity_code, "Unknown")

            # Check if the vulnerability name matches the filters
            if vuln_name in vuln_filters:
                hostname = ip_metadata.get(ip_address, {}).get("hostname", "Unknown Hostname")
                public_ip = ip_metadata.get(ip_address, {}).get("public_ip", "Unknown Public IP")
                vulnerabilities.append((ip_address, hostname, public_ip, vuln_name, severity, description, plugin_output, solution))

    return vulnerabilities


def add_cover_page(doc, report_date):
    """Add a cover page to the DOCX document."""
    cover_title = "Vulnerability Report"
    description = "This document provides a detailed analysis of vulnerabilities identified during a penetration test."

    # Add title
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title_paragraph.add_run(cover_title)
    run.bold = True
    run.font.size = Pt(24)

    # Add empty lines
    doc.add_paragraph()
    doc.add_paragraph()

    # Add date
    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_paragraph.add_run(f"Date: {report_date}").font.size = Pt(14)

    # Add empty lines
    doc.add_paragraph()
    doc.add_paragraph()

    # Add description
    description_paragraph = doc.add_paragraph()
    description_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    description_paragraph.add_run(description).font.size = Pt(12)

    # Add page break
    doc.add_page_break()


def save_to_docx(vulnerabilities, output_filename):
    # Create a new DOCX document
    doc = Document()

    # Generate report date
    report_date = datetime.now().strftime("%B %d, %Y")

    # Add cover page
    add_cover_page(doc, report_date)

    # Add title
    title = doc.add_heading("Vulnerability Report", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add Table of Contents
    doc.add_heading("Table of Contents", level=2)
    doc.add_paragraph("1. Scope", style="List Number")
    doc.add_paragraph("2. Summary of Findings", style="List Number")
    doc.add_paragraph("3. Details of Findings", style="List Number")

    # Add Scope
    doc.add_heading("1. Scope", level=2)
    doc.add_paragraph("This report summarizes the vulnerabilities identified during a penetration test of the target systems.")

    # Add Summary of Findings
    doc.add_heading("2. Summary of Findings", level=2)

    # Create summary
    df = pd.DataFrame(vulnerabilities, columns=["IP Address", "Hostname", "Public IP", "Vulnerability Name", "Severity", "Description", "Plugin Output", "Solution"])
    severity_counts = df["Severity"].value_counts().to_dict()

    # Add severity table
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Severity"
    hdr_cells[1].text = "Count"

    for severity, count in severity_counts.items():
        row_cells = table.add_row().cells
        row_cells[0].text = severity
        row_cells[1].text = str(count)

    # Add Details of Findings
    doc.add_heading("3. Details of Findings", level=2)

    # Group vulnerabilities by Vulnerability Name
    findings_by_vuln_name = defaultdict(list)
    for ip_address, hostname, public_ip, vuln_name, severity, description, plugin_output, solution in vulnerabilities:
        findings_by_vuln_name[vuln_name].append((ip_address, hostname, public_ip, severity, description, plugin_output, solution))

    # Add findings grouped by Vulnerability Name
    for index, (vuln_name, findings) in enumerate(sorted(findings_by_vuln_name.items()), start=1):
        # Add numbered Vulnerability Name Header
        doc.add_heading(f"{index}. Vulnerability Name: {vuln_name}", level=3)

        # Get unique severity and solution (assuming they're consistent across all IPs for a given vuln)
        severity = findings[0][3]
        description = findings[0][4]
        solution = findings[0][6]

        # Add Severity and Description
        doc.add_paragraph(f"Severity: {severity}", style="Normal")
        doc.add_paragraph(f"Description: {description}", style="Normal")

        # List affected IP Addresses
        doc.add_paragraph("Affected IP Addresses:", style="Normal")
        for ip_address, hostname, public_ip, _, _, plugin_output, _ in findings:
            doc.add_paragraph(f"- {ip_address} ({hostname}, Public IP: {public_ip})", style="List Bullet")

        # Add Plugin Output
        doc.add_paragraph("Plugin Output:", style="Normal")
        for ip_address, hostname, public_ip, _, _, plugin_output, _ in findings:
            create_shaded_table(doc, f"{ip_address}: {plugin_output}")

        # Add Solution
        doc.add_paragraph("Solution:", style="Normal")
        doc.add_paragraph(solution, style="Normal")

        # Add separator
        doc.add_paragraph("---", style="Normal")

    # Save the document
    doc.save(output_filename)
    print(f"Saved DOCX file to: {output_filename}")


def main():
    # Set up argument parser
    parser = argparse.ArgumentParser(description="Process a .nessus file and extract specific vulnerabilities.")
    parser.add_argument("nessus_file", help="Path to the .nessus file to be processed")
    parser.add_argument("vuln_file", help="Path to the file containing a list of vulnerabilities to include in the report")
    parser.add_argument("ip_filter", help="Path to the file containing IP address and metadata for filtering")
    args = parser.parse_args()

    # Parse IP and vulnerability filters
    vuln_filters = parse_vuln_file(args.vuln_file)
    ip_metadata = parse_ip_file(args.ip_filter) if args.ip_filter else {}
    ip_filter = set(ip_metadata.keys())

    # Parse the .nessus file for the specified vulnerability names
    vulnerabilities = parse_nessus_file_for_vulns(args.nessus_file, vuln_filters, ip_filter, ip_metadata)

    if not vulnerabilities:
        print("No matching vulnerabilities found.")
        return

    # Generate timestamp for filenames
    timestamp = datetime.now().strftime("%Y%m%d%H%M")

    # Define filename
    docx_filename = f"vulnerability_report_{timestamp}.docx"

    # Save results to DOCX
    save_to_docx(vulnerabilities, docx_filename)


if __name__ == '__main__':
    main()
