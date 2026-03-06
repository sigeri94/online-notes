import csv
from collections import defaultdict
from lxml import etree
from docx import Document
import matplotlib.pyplot as plt
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
import os


# ==============================
# Load Vulnerability Database
# ==============================

def load_vuln_db():

    vuln_db = {}

    with open("vuln_db.csv", newline='', encoding="utf-8") as f:

        reader = csv.DictReader(f)

        for row in reader:
            vuln_db[row["name"].lower()] = row

    return vuln_db


# ==============================
# Load Collapse Rules
# ==============================

def load_rules(file):

    rules = []

    with open(file, encoding="utf-8") as f:

        reader = csv.DictReader(f)

        for row in reader:

            rules.append(row)

    return rules


# ==============================
# Get vulnerability info safely
# ==============================

def get_vuln_info(finding, vuln_db):

    key = finding.lower()

    if key in vuln_db:
        return vuln_db[key]

    # fallback
    return {
        "severity": "Info",
        "cvss": "N/A",
        "description": "This vulnerability was detected by the scanning tool but is not yet documented in the local vulnerability database.",
        "impact": "Potential security risk depending on the configuration of the affected service.",
        "recommendation": "Review the service configuration and apply vendor recommended security patches."
    }


# ==============================
# Parse Nessus File
# ==============================

def parse_nessus(file):

    tree = etree.parse(file)

    vulns = []

    for host in tree.xpath("//ReportHost"):

        ip = host.attrib.get("name")

        for item in host.xpath(".//ReportItem"):

            plugin = item.attrib.get("pluginName")
            severity = item.attrib.get("severity")

            vulns.append({
                "ip": ip,
                "plugin": plugin,
                "severity": severity
            })

    return vulns


# ==============================
# Collapse Findings
# ==============================

def collapse_findings(vulns, rules):

    findings = defaultdict(set)

    for v in vulns:

        plugin = v["plugin"]

        for rule in rules:

            if rule["keyword"].lower() in plugin.lower():

                findings[rule["finding"]].add(v["ip"])

    return findings


# ==============================
# Generate Severity Chart
# ==============================

def generate_chart(findings, vuln_db):

    severity_count = defaultdict(int)

    for finding in findings:

        data = get_vuln_info(finding, vuln_db)

        sev = data["severity"]

        severity_count[sev] += 1

    labels = list(severity_count.keys())
    sizes = list(severity_count.values())

    plt.figure()

    plt.pie(sizes, labels=labels, autopct='%1.1f%%')

    plt.title("Vulnerability Severity Distribution")

    plt.savefig("output/severity_chart.png")

    plt.close()


# ==============================
# Generate Formal DOCX Report
# ==============================

def generate_formal_report(findings, vuln_db):

    doc = Document()

    # ======================
    # COVER PAGE
    # ======================

    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run = title.add_run("Penetration Test Report")

    run.bold = True
    run.font.size = Pt(30)

    doc.add_paragraph("")
    doc.add_paragraph("")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.add_run("Confidential Security Assessment")

    doc.add_paragraph("")
    doc.add_paragraph("")

    date = doc.add_paragraph()
    date.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date.add_run(datetime.now().strftime("%B %Y"))

    doc.add_page_break()

    # ======================
    # DOCUMENT CONTROL
    # ======================

    doc.add_heading("Document Control", level=1)

    table = doc.add_table(rows=1, cols=4)

    hdr = table.rows[0].cells

    hdr[0].text = "Version"
    hdr[1].text = "Date"
    hdr[2].text = "Author"
    hdr[3].text = "Description"

    row = table.add_row().cells

    row[0].text = "1.0"
    row[1].text = datetime.now().strftime("%Y-%m-%d")
    row[2].text = "Security Team"
    row[3].text = "Initial Release"

    doc.add_page_break()

    # ======================
    # EXECUTIVE SUMMARY
    # ======================

    doc.add_heading("Executive Summary", level=1)

    doc.add_paragraph(
        "This report presents the results of a security assessment conducted "
        "against the target infrastructure. The objective of this engagement "
        "was to identify vulnerabilities that may expose the environment to "
        "security threats."
    )

    # ======================
    # METHODOLOGY
    # ======================

    doc.add_heading("Methodology", level=1)

    doc.add_paragraph(
        "Testing activities were conducted using automated vulnerability "
        "scanning tools combined with manual verification techniques "
        "following industry standards such as PTES and OWASP Testing Guide."
    )

    # ======================
    # RISK SUMMARY
    # ======================

    doc.add_heading("Risk Rating Summary", level=1)

    severity_count = defaultdict(int)

    for finding in findings:

        data = get_vuln_info(finding, vuln_db)

        severity_count[data["severity"]] += 1

    table = doc.add_table(rows=1, cols=2)

    hdr = table.rows[0].cells
    hdr[0].text = "Severity"
    hdr[1].text = "Number of Findings"

    for sev, count in severity_count.items():

        row = table.add_row().cells

        row[0].text = sev
        row[1].text = str(count)

    doc.add_picture("output/severity_chart.png")

    doc.add_page_break()

    # ======================
    # SUMMARY TABLE
    # ======================

    doc.add_heading("Summary of Findings", level=1)

    table = doc.add_table(rows=1, cols=3)

    hdr = table.rows[0].cells
    hdr[0].text = "Finding"
    hdr[1].text = "Severity"
    hdr[2].text = "Affected Hosts"

    for finding, ips in findings.items():

        data = get_vuln_info(finding, vuln_db)

        row = table.add_row().cells

        row[0].text = finding
        row[1].text = data["severity"]
        row[2].text = str(len(ips))

    doc.add_page_break()

    # ======================
    # DETAILED FINDINGS
    # ======================

    doc.add_heading("Detailed Findings", level=1)

    index = 1

    for finding, ips in findings.items():

        data = get_vuln_info(finding, vuln_db)

        doc.add_heading(f"{index}. {finding}", level=2)

        doc.add_paragraph(f"Severity: {data['severity']}")
        doc.add_paragraph(f"CVSS Score: {data.get('cvss','N/A')}")

        doc.add_heading("Description", level=3)
        doc.add_paragraph(data["description"])

        doc.add_heading("Impact", level=3)
        doc.add_paragraph(data["impact"])

        doc.add_heading("Affected Systems", level=3)

        for ip in ips:
            doc.add_paragraph(ip)

        doc.add_heading("Recommendation", level=3)
        doc.add_paragraph(data["recommendation"])

        doc.add_page_break()

        index += 1

    doc.save("output/formal_pentest_report.docx")


# ==============================
# Main
# ==============================

def main():

    os.makedirs("output", exist_ok=True)

    print("[*] Parsing Nessus scan...")
    vulns = parse_nessus("scan.nessus")

    print("[*] Loading collapse rules...")
    rules = load_rules("collapse_rules.csv")

    print("[*] Loading vulnerability database...")
    vuln_db = load_vuln_db()

    print("[*] Collapsing findings...")
    findings = collapse_findings(vulns, rules)

    print("[*] Generating severity chart...")
    generate_chart(findings, vuln_db)

    print("[*] Generating report...")
    generate_formal_report(findings, vuln_db)

    print("[+] Report generated: output/formal_pentest_report.docx")


if __name__ == "__main__":
    main()